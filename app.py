# app.py
import os
import uuid
import json
import io
import numpy as np
import pandas as pd
from flask import Flask, request, jsonify, render_template
from docx import Document
from werkzeug.utils import secure_filename

app = Flask(__name__)
data_store = {}

# ------------------------------------------------------------
# чтение файла в DataFrame
# ------------------------------------------------------------
def read_file_to_df(file, ext):
    """
    file: файл‑объект (BytesIO) или путь к файлу
    ext: расширение (csv, txt, json, docx)
    """
    if ext == 'csv':
        return pd.read_csv(file, encoding='utf-8')
    elif ext == 'txt':
        content = file.read().decode('utf-8')
        return pd.read_csv(io.StringIO(content), sep=',')
    elif ext == 'json':
        return pd.read_json(io.TextIOWrapper(file, encoding='utf-8'))
    elif ext == 'docx':
        doc = Document(io.BytesIO(file.read()))
        if not doc.tables:
            raise ValueError("Документ не содержит таблиц")
        table = doc.tables[0]
        data = [[cell.text for cell in row.cells] for row in table.rows]
        df = pd.DataFrame(data[1:], columns=data[0])
        for col in df.columns:
            df[col] = pd.to_numeric(df[col], errors='ignore')
        return df
    else:
        raise ValueError("Неподдерживаемый формат")
# ------------------------------------------------------------
# маршруты
# ------------------------------------------------------------
@app.route('/')
def index():
    """Главная страница – рендеринг шаблона"""
    return render_template('dashboard.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'Нет файла'}), 400

    file = request.files['file']
    filename = secure_filename(file.filename)
    ext = filename.rsplit('.', 1)[1].lower() if '.' in filename else ''

    file_content = io.BytesIO(file.read())

    try:
        df = read_file_to_df(file_content, ext)
        file_key = uuid.uuid4().hex
        data_store[file_key] = df

        columns = df.columns.tolist()
        return jsonify({'columns': columns, 'file_key': file_key})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/analyze', methods=['POST'])
def analyze_data():
    data = request.json
    file_key = data.get('file_key')
    x_col = data.get('x_col')
    y_col = data.get('y_col')

    if not file_key or file_key not in data_store:
        return jsonify({'error': 'Файл не найден'}), 400
    if not x_col or not y_col:
        return jsonify({'error': 'Выберите колонки'}), 400

    try:
        df = data_store[file_key]

        if x_col not in df.columns or y_col not in df.columns:
            raise ValueError("Неверные колонки")

        # конвертация Y в числа
        df[y_col] = pd.to_numeric(df[y_col], errors='coerce')
        if df[y_col].isnull().all():
            raise ValueError("Y‑колонка не числовая")

        try:
            df[x_col] = pd.to_datetime(df[x_col])
            labels = df[x_col].dt.strftime('%Y-%m-%d').tolist()
        except Exception:
            labels = df[x_col].astype(str).tolist()

        avg = float(df[y_col].mean())
        mn = float(df[y_col].min())
        mx = float(df[y_col].max())
        pts = int(len(df))

        stats = {
            'average': avg,
            'min': mn,
            'max': mx,
            'data_points': pts
        }

        values = df[y_col].astype(float).tolist()
        chart_data = {
            'labels': labels,
            'values': values
        }

        return jsonify({
            'stats': stats,
            'chart_data': chart_data,
            'x_label': x_col,
            'y_label': y_col
        })
    except Exception as e:
        app.logger.exception("Error during analysis")
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)